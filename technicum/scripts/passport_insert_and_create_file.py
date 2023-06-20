import os
import datetime

from transliterate import translit

from docx import Document
from docx.shared import Inches
from django.conf import settings


def main_func(fields_dict: dict):
    '''1-берет шаблон документа
       2-скачивает текстовые поля и изображение из html страницы
       3-создает новые документ на основании шаблона
       4-вставляет в шаблон текст, изображения
       5-удаляет пустые поля в документе
       6-удаляет изображение из общей папки
       7-сохраняет готовый файл в общей папке
    '''
    if fields_dict['doc_type'][0:3] == 'smk':
        doc_type = fields_dict['doc_type']
        base_file_name = fields_dict['base_file_name']
        print('base_file_name =====', base_file_name)
        base_doc_path = os.path.join(settings.MEDIA_ROOT, 'smk', f'{base_file_name}')
        document = Document(base_doc_path)
        insert_data(fields_dict, document)  # вставляем данные в шаблонный документ
        new_doc_path = os.path.join(settings.MEDIA_ROOT, 'new_smk_files', f'new_{doc_type}.docx')
    else:
        doc_type = fields_dict['doc_type']
        base_doc_path = os.path.join(settings.MEDIA_ROOT, f'{doc_type}_base.docx')
        document = Document(base_doc_path)
        insert_data(fields_dict, document)  # вставляем данные в шаблонный документ
        new_doc_path = os.path.join(settings.MEDIA_ROOT, f'new_{doc_type}.docx')
    document.save(new_doc_path)
    count_execution()


def delete_file(file_path) -> None:
    '''Удаляет файл по адресу file_path'''
    if os.path.exists(file_path):
        os.remove(file_path)


def insert_data(fields_dict: dict, document) -> None:
    '''Вставляем изображение, текст в document. Удаляем пустые поля.'''
    if 'image_file_path' in fields_dict.keys():
        try:  # Если пользователь загрузил изображение jpeg, GIF, png
            insert_image_data(fields_dict, document)
        except: # Если пользователь якорь и загрузил не изображение
            file_path_to_del = fields_dict['image_file_path']
            image_path = os.path.join(settings.MEDIA_ROOT, f'{file_path_to_del}')
            delete_file(image_path)
    insert_text_data(fields_dict, document)
    delete_pf_from_doc(fields_dict, document)


def insert_text_data(fields_dict: dict, document) -> None:
    '''Заменяем метки document на нужные строки'''
    for paragraph in document.paragraphs:
        for field in fields_dict.keys():
            if field in paragraph.text:
                paragraph.text = paragraph.text.replace('{' + field + '}',
                                                        fields_dict[field])
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for field in fields_dict.keys():
                        if field in paragraph.text:
                            paragraph.text = paragraph.text.replace('{' + field + '}',
                                                                    fields_dict[field])


def insert_image_data(fields_dict: dict, document) -> None:
    '''Заменяем метку pf18 в document на изображение'''
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{pf18}' in paragraph.text:
                        paragraph.text = 'Рисунок 1 – Общий вид изделия\n'
                        image_file_path = fields_dict['image_file_path']
                        #image_path = os.path.join(settings.MEDIA_ROOT,
                                                  #f'{image_name}'
                                                  #)
                        run = paragraph.add_run()
                        picture = run.add_picture(image_file_path)
                        # Задаем максимально допустимый размер
                        max_width = Inches(4)
                        max_height = Inches(3)
                        # Получаем текущий размер изображения
                        current_width = picture.width
                        current_height = picture.height
                        # Вычисляем коэффициенты масштабирования
                        width_ratio = max_width / current_width
                        height_ratio = max_height / current_height
                        scale_ratio = min(width_ratio, height_ratio)
                        # Масштабируем изображение
                        picture.width = int(current_width * scale_ratio)
                        picture.height = int(current_height * scale_ratio)
                        # Удаляем временное изображение
                        delete_file(image_file_path)
                        break


def delete_pf_from_doc(fields_dict, document) -> None:
    '''Удаляет пустые метки в document'''
    clean_dict = get_clean_dict(fields_dict)
    sign_list = get_sign_list()
    for paragraph in document.paragraphs:
        for key, value in clean_dict.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
        for sign in sign_list:
            if sign in paragraph.text:
                paragraph.text = paragraph.text.replace(sign, '')
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in clean_dict.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)
                    for sign in sign_list:
                        if sign in paragraph.text:
                            paragraph.text = paragraph.text.replace(sign, '')


def get_sign_list() -> list:
    '''Возвращает список меток в формате STR от pf1 до pf59'''
    signs = []
    for i in range(60):
        str = f'pf{i}'
        sign = f'{{{str}}}'
        signs.append(sign)
    return signs


def get_clean_dict(fields_dict: dict) -> dict:
    '''Возвращает словарь clean_dict с данными
       , которые нужно заменить в доументе'''
    clean_dict = {
            'Серии: {pf3}': '',
            'Изделие предназначено для: {pf19}.': '',
            ': {pf19}': '',
            'Адрес изготовителя: {pf7}.': '',
            'Телефон: {pf8}.': '',
            '..': '.',
            'ОКПД2 {pf6}': '',
            '{pf5}': '000000',
            'ИНН {pf10}': '',
            'INN {pf10}': '',
            'INN {pf44}': '',
            'ИНН {pf44}': '',
            'parties 1 года': 'parties 1 year',
            'parties 2 лет': 'parties 2 years',
            'parties 3 лет': 'parties 3 years',
            'parties 5 лет': 'parties 5 years',
            'parties 10 лет': 'parties 10 years',
            'parties 20 лет': 'parties 20 years',
        }
    try:
        clean_dict['{pf11}'] = transliterate_to_english(fields_dict['pf1'])
    except:
        pass
    try:
        clean_dict['{pf13}'] = transliterate_to_english(fields_dict['pf3'])
    except:
        pass
    try:
        clean_dict['{pf14}'] = transliterate_to_english(fields_dict['pf4'])
    except:
        pass
    try:
        clean_dict['{pf15}'] = transliterate_to_english(fields_dict['pf6'])
    except:
        pass
    try:
        clean_dict['{pf46}'] = str((int(fields_dict['pf21']))*(int(fields_dict['pf22'])))
    except:
        pass
    try:
        clean_dict['{pf47}'] = str((int(fields_dict['pf24']))*(int(fields_dict['pf25'])))
    except:
        pass
    try:
        clean_dict['{pf48}'] = str((int(fields_dict['pf27']))*(int(fields_dict['pf28'])))
    except:
        pass
    try:
        clean_dict['{pf49}'] = str((int(fields_dict['pf30']))*(int(fields_dict['pf31'])))
    except:
        pass
    try:
        clean_dict['{pf50}'] = str((int(fields_dict['pf33']))*(int(fields_dict['pf34'])))
    except:
        pass
    try:
        clean_dict['{pf51}'] = str((int(fields_dict['pf36']))*(int(fields_dict['pf37'])))
    except:
        pass
    try:
        clean_dict['{pf52}'] = str((int(fields_dict['pf39']))*(int(fields_dict['pf40'])))
    except:
        pass
    try:
        clean_dict['{pf53}'] = str((int(fields_dict['pf42']))*(int(fields_dict['pf43'])))
    except:
        pass
    summ_amount = 0
    for ind in range(46, 54):
        keytext = f'pf{ind}'
        try:
            summ_amount = summ_amount + int(clean_dict['{'+keytext+'}'])
        except:
            pass
    clean_dict['{pf54}'] = str(summ_amount)

    return clean_dict


def transliterate_to_english(word: str) -> str:
    '''Переводит (транслитом) слово word'''
    transliterated_word = translit(word, 'ru', reversed=True)
    return transliterated_word


#   def count_execution() -> None:
#       '''Сохраняет количество обращений к функции
#       в файл counter.txt в корне проекта
#       '''
#       counter_path = os.path.join(settings.BASE_DIR, 'counter.txt')
#       try:
#           with open(counter_path, 'r') as file:
#               count = int(file.read())
#       except:
#           count = 0
#       count += 1
#       with open(counter_path, 'w') as file:
#           file.write(str(count))

def count_execution() -> None:
    '''Сохраняет количество обращений к функции
       в файл counter_date.txt в папке counter
    '''
    current_date = datetime.datetime.now().strftime("%Y-%m-%d")
    # Создаем папку "counter", если она не существует
    counter_dir = os.path.join(settings.BASE_DIR, 'counter')
    os.makedirs(counter_dir, exist_ok=True)
    counter_path = os.path.join(counter_dir, f'counter_{current_date}.txt')
    try:
        with open(counter_path, 'r') as file:
            count = int(file.read())
    except FileNotFoundError:
        count = 0
    count += 1
    with open(counter_path, 'w') as file:
        file.write(str(count))
